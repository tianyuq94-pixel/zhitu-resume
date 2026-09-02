from io import BytesIO

import pymupdf
import pytest
from docx import Document

from app.parsers.resume import DOCX_MIME, PDF_MIME, ResumeParseError, parse_resume
from app.storage.local import LocalResumeStorage


def make_pdf(text: str) -> bytes:
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), text)
    content = document.tobytes()
    document.close()
    return content


def make_docx(text: str) -> bytes:
    document = Document()
    document.add_heading("Resume", level=1)
    document.add_paragraph(text)
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_pdf_resume_text_is_extracted() -> None:
    parsed = parse_resume(
        make_pdf("Education: Example University. Experience: Built a production web application."),
        "resume.pdf",
        PDF_MIME,
    )

    assert parsed.mime_type == PDF_MIME
    assert "Example University" in parsed.text


def test_docx_resume_text_is_extracted() -> None:
    parsed = parse_resume(
        make_docx("Education: Example University. Project: AI career assistant website."),
        "resume.docx",
        DOCX_MIME,
    )

    assert parsed.mime_type == DOCX_MIME
    assert "AI career assistant" in parsed.text


def test_blank_pdf_is_rejected() -> None:
    document = pymupdf.open()
    document.new_page()
    content = document.tobytes()
    document.close()

    with pytest.raises(ResumeParseError, match="未识别到足够文字"):
        parse_resume(content, "blank.pdf", PDF_MIME)


def test_storage_key_cannot_escape_root(tmp_path) -> None:
    storage = LocalResumeStorage(tmp_path)

    with pytest.raises(ValueError, match="Invalid storage key"):
        storage.path_for("../outside.pdf")

