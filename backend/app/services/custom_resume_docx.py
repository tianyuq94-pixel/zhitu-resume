from io import BytesIO

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from app.models.ai import CustomResume
from app.services.custom_resume_photo import get_custom_resume_photo_storage

FONT_NAME = "Microsoft YaHei"
PAGE_WIDTH_MM = 210
PAGE_HEIGHT_MM = 297
HORIZONTAL_MARGIN_MM = 15
CONTENT_WIDTH_MM = PAGE_WIDTH_MM - 2 * HORIZONTAL_MARGIN_MM


def _set_east_asian_font(element, font_name: str = FONT_NAME) -> None:
    run_properties = element.get_or_add_rPr()
    run_fonts = run_properties.find(qn("w:rFonts"))
    if run_fonts is None:
        run_fonts = OxmlElement("w:rFonts")
        run_properties.insert(0, run_fonts)
    for attribute in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        run_fonts.set(qn(attribute), font_name)


def _configure_style(style, *, size: float, bold: bool = False) -> None:
    style.font.name = FONT_NAME
    style.font.size = Pt(size)
    style.font.bold = bold
    _set_east_asian_font(style.element)


def _get_or_add_paragraph_style(document: DocumentObject, name: str, base: str = "Normal"):
    if name in document.styles:
        return document.styles[name]
    style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = document.styles[base]
    return style


def _set_cell_margins(cell: _Cell, *, top: int = 0, start: int = 0, bottom: int = 0, end: int = 0) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    margins = cell_properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        cell_properties.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table: Table, widths_mm: list[float]) -> None:
    widths_dxa = [round(width * 56.6929134) for width in widths_mm]
    total_width = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    table_properties = table._tbl.tblPr
    table_width = table_properties.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.insert(0, table_width)
    table_width.set(qn("w:w"), str(total_width))
    table_width.set(qn("w:type"), "dxa")

    table_indent = table_properties.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_properties.append(table_indent)
    table_indent.set(qn("w:w"), "0")
    table_indent.set(qn("w:type"), "dxa")

    layout = table_properties.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    borders = table_properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_properties.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "nil")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width))
        grid.append(grid_column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            cell_width.set(qn("w:w"), str(widths_dxa[index]))
            cell_width.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            _set_cell_margins(cell)


def _set_bottom_border(paragraph: Paragraph) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    paragraph_borders = paragraph_properties.find(qn("w:pBdr"))
    if paragraph_borders is None:
        paragraph_borders = OxmlElement("w:pBdr")
        paragraph_properties.append(paragraph_borders)
    bottom = paragraph_borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        paragraph_borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "000000")


def _add_square_bullet_numbering(document: DocumentObject) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level_type = OxmlElement("w:multiLevelType")
    multi_level_type.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level_type)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "bullet")
    level.append(number_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "▪")
    level.append(level_text)
    level_alignment = OxmlElement("w:lvlJc")
    level_alignment.set(qn("w:val"), "left")
    level.append(level_alignment)
    paragraph_properties = OxmlElement("w:pPr")
    indentation = OxmlElement("w:ind")
    indentation.set(qn("w:left"), "238")
    indentation.set(qn("w:hanging"), "187")
    paragraph_properties.append(indentation)
    level.append(paragraph_properties)
    run_properties = OxmlElement("w:rPr")
    run_fonts = OxmlElement("w:rFonts")
    for attribute in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        run_fonts.set(qn(attribute), FONT_NAME)
    run_properties.append(run_fonts)
    level.append(run_properties)
    abstract.append(level)
    numbering.append(abstract)

    num_ids = [int(element.get(qn("w:numId"))) for element in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    number.append(abstract_reference)
    numbering.append(number)
    return num_id


def _apply_numbering(paragraph: Paragraph, num_id: int) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    number_properties = paragraph_properties.find(qn("w:numPr"))
    if number_properties is None:
        number_properties = OxmlElement("w:numPr")
        paragraph_properties.append(number_properties)
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    number_properties.append(level)
    number_properties.append(number)


def _configure_document(document: DocumentObject) -> None:
    section = document.sections[0]
    section.page_width = Mm(PAGE_WIDTH_MM)
    section.page_height = Mm(PAGE_HEIGHT_MM)
    section.left_margin = Mm(HORIZONTAL_MARGIN_MM)
    section.right_margin = Mm(HORIZONTAL_MARGIN_MM)
    section.top_margin = Mm(14)
    section.bottom_margin = Mm(14)
    section.footer_distance = Mm(6.5)

    normal = document.styles["Normal"]
    _configure_style(normal, size=9.6)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = Pt(15.2)
    normal.paragraph_format.widow_control = True

    name = _get_or_add_paragraph_style(document, "ResumeName")
    _configure_style(name, size=21, bold=True)
    name.paragraph_format.space_before = Pt(0)
    name.paragraph_format.space_after = Mm(3.5)
    name.paragraph_format.line_spacing = Pt(27)

    contact = _get_or_add_paragraph_style(document, "ResumeContact")
    _configure_style(contact, size=9.4)
    contact.paragraph_format.space_before = Mm(1.4)
    contact.paragraph_format.space_after = Mm(1.4)
    contact.paragraph_format.line_spacing = Pt(14)

    section_style = _get_or_add_paragraph_style(document, "ResumeSection")
    _configure_style(section_style, size=13, bold=True)
    section_style.paragraph_format.space_before = Mm(3)
    section_style.paragraph_format.space_after = Mm(1.7)
    section_style.paragraph_format.line_spacing = Pt(17)
    section_style.paragraph_format.keep_with_next = True

    heading = _get_or_add_paragraph_style(document, "ResumeEntryHeading")
    _configure_style(heading, size=10.2, bold=True)
    heading.paragraph_format.space_before = Mm(1.8)
    heading.paragraph_format.space_after = Mm(0.8)
    heading.paragraph_format.line_spacing = Pt(15.5)
    heading.paragraph_format.keep_with_next = True

    bullet = _get_or_add_paragraph_style(document, "ResumeBullet")
    _configure_style(bullet, size=9.6)
    bullet.paragraph_format.left_indent = Mm(4.2)
    bullet.paragraph_format.first_line_indent = Mm(-3.3)
    bullet.paragraph_format.space_before = Pt(0)
    bullet.paragraph_format.space_after = Mm(1.15)
    bullet.paragraph_format.line_spacing = Pt(15.2)
    bullet.paragraph_format.widow_control = True


def _add_contact_table(container: _Cell, header: dict, width_mm: float) -> None:
    contacts = [
        ("联系电话", header.get("phone")),
        ("电子邮箱", header.get("email")),
        ("所在地", header.get("location")),
        ("出生年月", header.get("birth_date")),
    ]
    contacts = [(label, str(value).strip()) for label, value in contacts if str(value or "").strip()]
    if not contacts:
        return
    row_count = (len(contacts) + 1) // 2
    table = container.add_table(rows=row_count, cols=2)
    _set_table_geometry(table, [width_mm / 2, width_mm / 2])
    for index, (label, value) in enumerate(contacts):
        paragraph = table.cell(index // 2, index % 2).paragraphs[0]
        paragraph.style = "ResumeContact"
        paragraph.paragraph_format.right_indent = Mm(3)
        label_run = paragraph.add_run(f"{label}：")
        label_run.bold = True
        paragraph.add_run(value)


def _add_header(document: DocumentObject, header: dict) -> None:
    photo_key = header.get("_photo_storage_key")
    photo_path = get_custom_resume_photo_storage().path_for(photo_key) if photo_key else None
    has_photo = bool(photo_path and photo_path.is_file())
    widths = [151.0, 4.0, 25.0] if has_photo else [CONTENT_WIDTH_MM]
    table = document.add_table(rows=1, cols=len(widths))
    _set_table_geometry(table, widths)

    identity_cell = table.cell(0, 0)
    name_paragraph = identity_cell.paragraphs[0]
    name_paragraph.style = "ResumeName"
    name_paragraph.add_run(str(header.get("name") or "").strip() or "姓名")
    political_status = str(header.get("political_status") or "").strip()
    if political_status:
        status_run = name_paragraph.add_run(f"（{political_status}）")
        status_run.font.size = Pt(11)
        _set_east_asian_font(status_run._element)
    _add_contact_table(identity_cell, header, widths[0])

    if has_photo and photo_path is not None:
        photo_paragraph = table.cell(0, 2).paragraphs[0]
        photo_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        photo_paragraph.paragraph_format.space_after = Pt(0)
        photo_paragraph.paragraph_format.line_spacing = 1
        photo_paragraph.add_run().add_picture(str(photo_path), width=Mm(25), height=Mm(34))

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Mm(2)
    spacer.paragraph_format.line_spacing = Pt(1)


def _add_footer(document: DocumentObject) -> None:
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    prefix = paragraph.add_run("第 ")
    prefix.font.size = Pt(7.5)
    prefix.font.color.rgb = None
    field_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, display, end):
        field_run._r.append(element)
    suffix = paragraph.add_run(" 页")
    suffix.font.size = Pt(7.5)
    for run in paragraph.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(7.5)
        _set_east_asian_font(run._element)


def build_custom_resume_docx(custom_resume: CustomResume) -> bytes:
    document = Document()
    _configure_document(document)
    content = custom_resume.content or {}
    header = content.get("header") or {}
    name = str(header.get("name") or "").strip()
    document.core_properties.title = f"{name or '个人'}简历"
    document.core_properties.author = "职途简历"
    document.core_properties.subject = "个人简历"

    _add_header(document, header)
    bullet_num_id = _add_square_bullet_numbering(document)
    for section in content.get("sections", []):
        section_title = str(section.get("title") or "").strip()
        visible_items = [item for item in section.get("items", []) if str(item.get("final_text") or "").strip()]
        if not section_title or not visible_items:
            continue
        section_paragraph = document.add_paragraph(section_title, style="ResumeSection")
        _set_bottom_border(section_paragraph)
        for index, item in enumerate(visible_items):
            text = str(item.get("final_text") or "").strip()
            if item.get("item_type") == "heading":
                paragraph = document.add_paragraph(text, style="ResumeEntryHeading")
                paragraph.paragraph_format.keep_with_next = index < len(visible_items) - 1
            else:
                paragraph = document.add_paragraph(text, style="ResumeBullet")
                _apply_numbering(paragraph, bullet_num_id)

    _add_footer(document)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
